import os
import re
import socket
import io
from urllib.parse import urlparse
from fastapi import APIRouter, HTTPException
from fastapi.responses import HTMLResponse, FileResponse, StreamingResponse
from sqlalchemy import select
from jinja2 import Environment, FileSystemLoader
from datetime import datetime, timezone
from pydantic import BaseModel
from typing import Optional

from database import async_session, Project, Scan, Finding
from config import TEMPLATE_DIR, REPORTS_DIR

router = APIRouter()
jinja_env = Environment(loader=FileSystemLoader(TEMPLATE_DIR))

TOOL_PURPOSES = {
    "nmap": "Port scanning & service detection",
    "sqlmap": "SQL injection testing",
    "nikto": "Web server vulnerability scanning",
    "dirb": "Directory brute-forcing",
    "gobuster_dir": "Directory & file discovery",
    "ffuf": "Web fuzzing",
    "whatweb": "Web technology fingerprinting",
    "wfuzz": "Web application fuzzing",
    "nxc": "Network service enumeration & exploitation",
    "enum4linux": "SMB/Samba enumeration",
    "smbclient": "SMB share access",
    "searchsploit": "Exploit database search",
    "hydra": "Online brute-force attacks",
    "john": "Offline password cracking",
    "hashcat": "GPU-accelerated password cracking",
    "hashid": "Hash type identification",
    "amass": "Subdomain enumeration",
    "gobuster_dns": "DNS subdomain brute-force",
    "dnsenum": "DNS enumeration",
    "nuclei": "Template-based vulnerability scanning",
    "wafw00f": "WAF detection",
    "feroxbuster": "Content discovery",
    "wpscan": "WordPress security scanning",
    "whois": "Domain registration lookup",
    "dig": "DNS query tool",
}

REMEDIATION_MAP = {
    "open_port": "Review if this port/service is required. Disable unnecessary services and apply firewall rules.",
    "vuln_detected": "Apply vendor patches or upgrade to the latest version. Implement compensating controls if patching is not immediately possible.",
    "sql_injection": "Use parameterized queries/prepared statements. Implement input validation and WAF rules.",
    "credential_found": "Enforce strong password policies. Implement account lockout and MFA. Rotate compromised credentials immediately.",
    "exploit_available": "Patch or upgrade the affected software immediately. Monitor for exploitation attempts.",
    "default": "Investigate and remediate according to organizational security policy.",
}


def _resolve_ip(host: str) -> str:
    """Resolve a hostname/URL to its IP address."""
    try:
        parsed = urlparse(host)
        hostname = parsed.hostname or host
        hostname = hostname.strip().rstrip(".")
        if not hostname:
            return "N/A"
        info = socket.getaddrinfo(hostname, None, socket.AF_UNSPEC, socket.SOCK_STREAM)
        if info:
            return info[0][4][0]
    except Exception:
        pass
    return "N/A"


def _collect_targets_with_ips(target: str, scans: list) -> list[dict]:
    """Build a list of unique targets with their resolved IPs."""
    seen = set()
    targets = []

    primary = target.strip()
    if primary and primary not in seen:
        seen.add(primary)
        targets.append({"url": primary, "ip": _resolve_ip(primary), "description": "Primary target"})

    for s in scans:
        cmd = s.get("command", "") if isinstance(s, dict) else (s.command or "")
        for token in cmd.split():
            if "://" in token or (re.match(r"^\d+\.\d+\.\d+\.\d+", token)):
                cleaned = token.strip("'\"")
                if cleaned and cleaned not in seen and not cleaned.startswith("-"):
                    seen.add(cleaned)
                    targets.append({"url": cleaned, "ip": _resolve_ip(cleaned), "description": "Scanned target"})

    return targets


def extract_findings_from_output(tool_name: str, command: str, output: str) -> list[dict]:
    findings = []
    if not output:
        return findings

    if tool_name == "nmap":
        for m in re.finditer(r"(\d+)/(tcp|udp)\s+(open)\s+(\S+)\s*(.*)", output):
            port, proto, state, service, version = m.groups()
            version = version.strip()
            sev = "info"
            title = f"Open Port {port}/{proto} - {service}"
            desc = f"Port {port}/{proto} is open running {service}"
            if version:
                desc += f" ({version})"
                title += f" ({version})"
            findings.append({
                "severity": sev, "title": title, "description": desc,
                "evidence": m.group(0).strip(), "component": f"{port}/{proto}",
                "tool": tool_name, "command": command, "remediation": REMEDIATION_MAP["open_port"],
            })
        for m in re.finditer(r"(.*VULNERABLE.*)", output, re.I):
            findings.append({
                "severity": "high", "title": "Vulnerability Detected (Nmap NSE)",
                "description": m.group(1).strip(),
                "evidence": _extract_context(output, m.start(), 300),
                "component": "", "tool": tool_name, "command": command,
                "remediation": REMEDIATION_MAP["vuln_detected"],
            })
        for m in re.finditer(r"(CVE-\d{4}-\d+)", output):
            findings.append({
                "severity": "high", "title": f"CVE Reference: {m.group(1)}",
                "description": f"Nmap scripts identified {m.group(1)}",
                "evidence": _extract_context(output, m.start(), 200),
                "component": "", "tool": tool_name, "command": command,
                "remediation": REMEDIATION_MAP["vuln_detected"],
            })

    elif tool_name == "sqlmap":
        if re.search(r"(is vulnerable|injectable)", output, re.I):
            param_match = re.search(r"Parameter:\s*(\S+)", output)
            param = param_match.group(1) if param_match else "unknown"
            findings.append({
                "severity": "critical",
                "title": f"SQL Injection - Parameter '{param}'",
                "description": f"SQLMap confirmed SQL injection vulnerability in parameter '{param}'",
                "evidence": output[:2000], "component": param,
                "tool": tool_name, "command": command,
                "remediation": REMEDIATION_MAP["sql_injection"],
            })
        for m in re.finditer(r"Type:\s*(.+?)(?:\n|$)", output):
            technique = m.group(1).strip()
            findings.append({
                "severity": "critical",
                "title": f"SQL Injection Technique: {technique}",
                "description": f"SQLMap identified injection using: {technique}",
                "evidence": _extract_context(output, m.start(), 300),
                "component": "", "tool": tool_name, "command": command,
                "remediation": REMEDIATION_MAP["sql_injection"],
            })

    elif tool_name == "nikto":
        for m in re.finditer(r"\+\s+(OSVDB-\d+):\s*(.*?)(?:\n|$)", output):
            ref = m.group(1)
            detail = m.group(2).strip()
            sev = "medium"
            if re.search(r"VULNERABLE|injection|XSS|remote\s+code|backdoor", detail, re.I):
                sev = "critical"
            elif re.search(r"header|cookie|disclosure|version", detail, re.I):
                sev = "low"
            findings.append({
                "severity": sev,
                "title": f"Nikto [{ref}]: {detail[:70]}",
                "description": detail,
                "evidence": m.group(0).strip(),
                "component": "", "tool": tool_name, "command": command,
                "remediation": REMEDIATION_MAP["vuln_detected"],
            })
        for m in re.finditer(r"\+\s+(/\S+.*?)(?:\n|$)", output):
            line = m.group(1).strip()
            if len(line) > 10 and "OSVDB" not in line:
                sev = "info"
                if re.search(r"directory listing|index of|default|enabled", line, re.I):
                    sev = "medium"
                elif re.search(r"X-Frame|X-Content|Strict-Transport|Content-Security", line, re.I):
                    sev = "low"
                findings.append({
                    "severity": sev,
                    "title": f"Nikto: {line[:80]}",
                    "description": line, "evidence": line,
                    "component": "", "tool": tool_name, "command": command,
                    "remediation": REMEDIATION_MAP["default"],
                })

    elif tool_name == "whatweb":
        for m in re.finditer(r"\+\s+(/\S+.*?)(?:\n|$)", output):
            line = m.group(1).strip()
            if len(line) > 10:
                findings.append({
                    "severity": "info", "title": f"Web Finding: {line[:80]}",
                    "description": line, "evidence": line,
                    "component": "", "tool": tool_name, "command": command,
                    "remediation": REMEDIATION_MAP["default"],
                })

    elif tool_name == "searchsploit":
        for m in re.finditer(r"(exploits/\S+)", output):
            path = m.group(1)
            title_match = re.search(r"(.+?)\s*\|\s*" + re.escape(path), output)
            title = title_match.group(1).strip() if title_match else path
            findings.append({
                "severity": "high", "title": f"Exploit Available: {title[:80]}",
                "description": f"Public exploit found at {path}",
                "evidence": _extract_context(output, m.start(), 200),
                "component": "", "tool": tool_name, "command": command,
                "remediation": REMEDIATION_MAP["exploit_available"],
            })

    elif tool_name == "nxc":
        for m in re.finditer(r"(Pwn3d!|STATUS_PASSWORD_EXPIRED|STATUS_LOGON_FAILURE.*admin|SUCCESS)", output, re.I):
            findings.append({
                "severity": "critical" if "Pwn3d" in m.group(0) else "high",
                "title": f"NXC: {m.group(0).strip()[:60]}",
                "description": _extract_context(output, m.start(), 200),
                "evidence": _extract_context(output, m.start(), 300),
                "component": "", "tool": tool_name, "command": command,
                "remediation": REMEDIATION_MAP["credential_found"],
            })
        for m in re.finditer(r"\[\+\]\s*(.*)", output):
            line = m.group(1).strip()
            if line and "Pwn3d" not in line:
                findings.append({
                    "severity": "info", "title": f"NXC Finding: {line[:80]}",
                    "description": line, "evidence": line,
                    "component": "", "tool": tool_name, "command": command,
                    "remediation": REMEDIATION_MAP["default"],
                })

    elif tool_name in ("hydra", "john", "hashcat", "hashcat_crack", "john_crack"):
        for m in re.finditer(r"\[(\d+)\]\[(\w+)\]\s*host:\s*(\S+)\s+login:\s*(\S+)\s+password:\s*(\S+)", output):
            port, svc, host, user, pw = m.groups()
            findings.append({
                "severity": "critical",
                "title": f"Credential Found: {user}@{host}:{port} ({svc})",
                "description": f"Valid credentials discovered: {user}:{pw} on {host}:{port} ({svc})",
                "evidence": m.group(0), "component": f"{host}:{port}",
                "tool": tool_name, "command": command,
                "remediation": REMEDIATION_MAP["credential_found"],
            })
        for m in re.finditer(r"(\S+):(\S+)", output):
            if "password" in output.lower() and "cracked" in output.lower():
                findings.append({
                    "severity": "high",
                    "title": f"Hash Cracked: {m.group(1)[:30]}",
                    "description": "Cracked hash value found", "evidence": m.group(0),
                    "component": "", "tool": tool_name, "command": command,
                    "remediation": REMEDIATION_MAP["credential_found"],
                })
                break

    elif tool_name == "nuclei":
        for m in re.finditer(r"\[(\w+)\]\s*\[([^\]]+)\]\s*\[([^\]]+)\]\s*(.*)", output):
            sev_raw = m.group(1).lower()
            template = m.group(2)
            proto = m.group(3)
            rest = m.group(4).strip()
            sev_map = {"critical": "critical", "high": "high", "medium": "medium", "low": "low", "info": "info"}
            sev = sev_map.get(sev_raw, "info")
            findings.append({
                "severity": sev, "title": f"Nuclei: {template}",
                "description": f"[{proto}] {rest}",
                "evidence": m.group(0), "component": rest[:60],
                "tool": tool_name, "command": command,
                "remediation": REMEDIATION_MAP["vuln_detected"],
            })

    elif tool_name == "wafw00f":
        for m in re.finditer(r"is behind\s+(.+?)(?:\s+WAF)?\.?\s*$", output, re.I | re.M):
            waf_name = re.sub(r"\x1b\[[0-9;]*m", "", m.group(1).strip())
            findings.append({
                "severity": "medium",
                "title": f"WAF Detected: {waf_name}",
                "description": f"The target is protected by {waf_name} WAF. Verify WAF rules cover OWASP Top 10.",
                "evidence": _extract_context(output, m.start(), 300),
                "component": "", "tool": tool_name, "command": command,
                "remediation": "Ensure WAF rules are up-to-date. Test for WAF bypass techniques.",
            })
        if re.search(r"No WAF detected|is not behind a WAF", output, re.I):
            findings.append({
                "severity": "critical",
                "title": "No WAF Detected",
                "description": "The target does not appear to be protected by a Web Application Firewall.",
                "evidence": output[:500],
                "component": "", "tool": tool_name, "command": command,
                "remediation": "Deploy a WAF to protect against common web attacks (SQLi, XSS, RFI).",
            })

    return findings


def _extract_context(text: str, pos: int, length: int) -> str:
    start = max(0, pos - 50)
    end = min(len(text), pos + length)
    return text[start:end].strip()


class ReportRequest(BaseModel):
    tester_name: str = ""
    classification: str = "CONFIDENTIAL"
    scope_notes: str = ""
    include_raw: bool = True


async def _gather_report_data(project_id: int, req: ReportRequest = ReportRequest()):
    """Shared helper: load project, scans, findings, resolve IPs, filter info."""
    async with async_session() as session:
        project = await session.get(Project, project_id)
        if not project:
            raise HTTPException(404, "Project not found")

        result = await session.execute(
            select(Scan).where(Scan.project_id == project_id).order_by(Scan.started_at)
        )
        scans = result.scalars().all()

        fr = await session.execute(
            select(Finding).join(Scan).where(Scan.project_id == project_id)
        )
        db_findings = fr.scalars().all()

    all_findings = []
    for f in db_findings:
        all_findings.append({
            "severity": f.severity, "title": f.title,
            "description": f.description, "evidence": f.evidence,
            "component": "", "tool": "", "command": "", "remediation": "",
        })

    from tool_runner import tasks as mem_tasks
    scan_data = []
    for s in scans:
        output = s.output or ""
        for tid, t in mem_tasks.items():
            if t.get("command") == s.command and t.get("tool_name") == s.tool_name:
                if len(t.get("output", "")) > len(output):
                    output = t["output"]
                break
        auto = extract_findings_from_output(s.tool_name, s.command, output)
        all_findings.extend(auto)

        duration = ""
        if s.started_at and s.finished_at:
            delta = s.finished_at - s.started_at
            duration = f"{int(delta.total_seconds())}s"
        scan_data.append({
            "tool_name": s.tool_name, "command": s.command,
            "status": s.status, "output": output,
            "started_at": s.started_at.isoformat() if s.started_at else "",
            "finished_at": s.finished_at.isoformat() if s.finished_at else "",
            "duration": duration,
        })

    seen = set()
    unique_findings = []
    for f in all_findings:
        key = f"{f['title']}|{f['severity']}"
        if key not in seen:
            seen.add(key)
            unique_findings.append(f)

    actionable_findings = [f for f in unique_findings if f["severity"] in ("critical", "high", "medium", "low")]
    actionable_findings.sort(key=lambda x: {"critical": 0, "high": 1, "medium": 2, "low": 3}.get(x["severity"], 4))

    severity_counts = {}
    for f in actionable_findings:
        severity_counts[f["severity"]] = severity_counts.get(f["severity"], 0) + 1

    targets_with_ips = _collect_targets_with_ips(project.target, scan_data)

    return {
        "project": project,
        "scan_data": scan_data,
        "findings": actionable_findings,
        "severity_counts": severity_counts,
        "targets_with_ips": targets_with_ips,
        "total_scans": len(scans),
        "total_findings": len(actionable_findings),
    }


async def _render_html(project_id: int, req: ReportRequest = ReportRequest()) -> str:
    data = await _gather_report_data(project_id, req)
    project = data["project"]

    template = jinja_env.get_template("template.html")
    html = template.render(
        project_name=project.name,
        target=project.target,
        scope=req.scope_notes or project.scope,
        tester_name=req.tester_name,
        classification=req.classification,
        generated_at=datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
        scans=data["scan_data"],
        findings=data["findings"],
        total_scans=data["total_scans"],
        total_findings=data["total_findings"],
        severity_counts=data["severity_counts"],
        tool_purposes=TOOL_PURPOSES,
        include_raw=req.include_raw,
        targets_with_ips=data["targets_with_ips"],
    )

    filename = f"report_{project.name.replace(' ', '_')}_{project_id}.html"
    filepath = os.path.join(REPORTS_DIR, filename)
    with open(filepath, "w") as f:
        f.write(html)

    return html


@router.post("/{project_id}")
async def generate_report(project_id: int, req: ReportRequest = ReportRequest()):
    html = await _render_html(project_id, req)
    return HTMLResponse(content=html)


@router.get("/{project_id}")
async def get_report(project_id: int):
    return await generate_report(project_id)


@router.get("/{project_id}/download")
async def download_report(project_id: int):
    async with async_session() as session:
        project = await session.get(Project, project_id)
        if not project:
            raise HTTPException(404, "Project not found")

    filename = f"report_{project.name.replace(' ', '_')}_{project_id}.html"
    filepath = os.path.join(REPORTS_DIR, filename)
    if not os.path.exists(filepath):
        await _render_html(project_id)
    return FileResponse(filepath, filename=filename, media_type="text/html")


@router.get("/{project_id}/pdf")
async def download_pdf(project_id: int):
    """Generate the HTML report and convert to A4 PDF via xhtml2pdf."""
    html = await _render_html(project_id)

    from xhtml2pdf import pisa

    pdf_buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(io.StringIO(html), dest=pdf_buffer)
    if pisa_status.err:
        raise HTTPException(500, "PDF generation failed")
    pdf_buffer.seek(0)

    async with async_session() as session:
        project = await session.get(Project, project_id)
    name = project.name.replace(" ", "_") if project else "report"
    filename = f"report_{name}_{project_id}.pdf"

    return StreamingResponse(
        pdf_buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{project_id}/docx")
async def download_docx(project_id: int):
    """Generate a Word document from report data."""
    data = await _gather_report_data(project_id)
    project = data["project"]

    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Cover page
    for _ in range(6):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SAWLAH")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Penetration Testing Framework")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Penetration Test Report")
    run.bold = True
    run.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(project.target)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph("")
    doc.add_paragraph("")

    meta_items = [
        ("Project", project.name),
        ("Date", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")),
        ("Tester", "Security Team"),
        ("Classification", "CONFIDENTIAL"),
        ("Total Findings", str(data["total_findings"])),
        ("Total Scans", str(data["total_scans"])),
    ]
    meta_table = doc.add_table(rows=len(meta_items), cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(meta_items):
        meta_table.cell(i, 0).text = label
        meta_table.cell(i, 1).text = value
        for cell in meta_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    doc.add_page_break()

    # 1. Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        f"A penetration test was conducted against {project.target} as part of the "
        f"{project.name} engagement. A total of {data['total_scans']} scans were executed, "
        f"identifying {data['total_findings']} actionable findings."
    )

    sev_table = doc.add_table(rows=2, cols=4)
    sev_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sev_labels = ["Critical", "High", "Medium", "Low"]
    sev_keys = ["critical", "high", "medium", "low"]
    sev_colors = [
        RGBColor(0xC0, 0x39, 0x2B),
        RGBColor(0xE6, 0x7E, 0x22),
        RGBColor(0xD4, 0xA0, 0x17),
        RGBColor(0x34, 0x98, 0xDB),
    ]
    for i, (label, key, color) in enumerate(zip(sev_labels, sev_keys, sev_colors)):
        count_cell = sev_table.cell(0, i)
        count_cell.text = str(data["severity_counts"].get(key, 0))
        for p in count_cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(20)
                run.font.color.rgb = color

        label_cell = sev_table.cell(1, i)
        label_cell.text = label
        for p in label_cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")

    # 2. Scope
    doc.add_heading("2. Scope and Targets", level=1)
    targets = data["targets_with_ips"]
    if targets:
        t_table = doc.add_table(rows=1 + len(targets), cols=3)
        t_table.style = "Table Grid"
        t_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        headers = ["URL / Host", "IP Address", "Description"]
        for i, h in enumerate(headers):
            t_table.cell(0, i).text = h
            for p in t_table.cell(0, i).paragraphs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(10)
        for i, t in enumerate(targets, 1):
            t_table.cell(i, 0).text = t["url"]
            t_table.cell(i, 1).text = t["ip"]
            t_table.cell(i, 2).text = t["description"]

    doc.add_paragraph("")

    # 3. Findings Summary
    doc.add_heading("3. Findings Summary", level=1)
    findings = data["findings"]
    if findings:
        f_table = doc.add_table(rows=1 + len(findings), cols=5)
        f_table.style = "Table Grid"
        f_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        f_headers = ["#", "Title", "Severity", "Component", "Tool"]
        for i, h in enumerate(f_headers):
            f_table.cell(0, i).text = h
            for p in f_table.cell(0, i).paragraphs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(9)

        for i, f in enumerate(findings, 1):
            f_table.cell(i, 0).text = f"F-{i:03d}"
            f_table.cell(i, 1).text = f["title"][:60]
            f_table.cell(i, 2).text = f["severity"].upper()
            f_table.cell(i, 3).text = (f["component"] or project.target)[:40]
            f_table.cell(i, 4).text = f["tool"] or "N/A"
            for cell in f_table.rows[i].cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
    else:
        doc.add_paragraph("No actionable findings were identified during this assessment.")

    doc.add_page_break()

    # 4. Detailed Findings
    doc.add_heading("4. Detailed Findings", level=1)
    for i, f in enumerate(findings, 1):
        doc.add_heading(f"F-{i:03d} [{f['severity'].upper()}] {f['title'][:70]}", level=2)
        doc.add_paragraph(f"Severity: {f['severity'].upper()}")
        doc.add_paragraph(f"Component: {f['component'] or project.target}")
        doc.add_paragraph(f"Tool: {f['tool'] or 'N/A'}")
        doc.add_paragraph("")
        doc.add_paragraph(f["description"])

        if f.get("evidence"):
            doc.add_paragraph("")
            p = doc.add_paragraph()
            run = p.add_run("Proof of Concept:")
            run.bold = True
            if f.get("command"):
                p_cmd = doc.add_paragraph()
                run_cmd = p_cmd.add_run(f"$ {f['command']}")
                run_cmd.font.name = "Consolas"
                run_cmd.font.size = Pt(9)
                run_cmd.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            p_ev = doc.add_paragraph()
            run_ev = p_ev.add_run(f["evidence"][:2000])
            run_ev.font.name = "Consolas"
            run_ev.font.size = Pt(8)

        if f.get("remediation"):
            doc.add_paragraph("")
            p = doc.add_paragraph()
            run = p.add_run("Remediation: ")
            run.bold = True
            p.add_run(f["remediation"])

        doc.add_paragraph("")

    # 5. Tool Execution Log
    doc.add_page_break()
    doc.add_heading("5. Tool Execution Log", level=1)
    scan_data = data["scan_data"]
    if scan_data:
        log_table = doc.add_table(rows=1 + len(scan_data), cols=4)
        log_table.style = "Table Grid"
        log_headers = ["Tool", "Command", "Status", "Duration"]
        for i, h in enumerate(log_headers):
            log_table.cell(0, i).text = h
            for p in log_table.cell(0, i).paragraphs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(9)
        for i, s in enumerate(scan_data, 1):
            log_table.cell(i, 0).text = s["tool_name"]
            log_table.cell(i, 1).text = s["command"][:60]
            log_table.cell(i, 2).text = s["status"]
            log_table.cell(i, 3).text = s["duration"] or "N/A"
            for cell in log_table.rows[i].cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    name = project.name.replace(" ", "_")
    filename = f"report_{name}_{project_id}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
